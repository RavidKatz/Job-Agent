import sys
from pathlib import Path


def read_text(path: Path) -> str:
    data = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "cp1255", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise RuntimeError("PDF extraction requires the pypdf package.") from exc

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def read_docx(path: Path) -> str:
    try:
        import docx
    except Exception as exc:
        raise RuntimeError("DOCX extraction requires the python-docx package.") from exc

    document = docx.Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = []
    for table in document.tables:
        for row in table.rows:
            table_cells.extend(cell.text for cell in row.cells)
    return "\n".join([*paragraphs, *table_cells])


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: extract_resume.py <resume-file>", file=sys.stderr)
        return 2

    path = Path(sys.argv[1])
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".csv"}:
        text = read_text(path)
    elif suffix == ".pdf":
        text = read_pdf(path)
    elif suffix == ".docx":
        text = read_docx(path)
    else:
        raise RuntimeError(f"Unsupported resume file type: {suffix}")

    sys.stdout.buffer.write(text.encode("utf-8", errors="replace"))
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(str(exc), file=sys.stderr)
        raise SystemExit(1)
